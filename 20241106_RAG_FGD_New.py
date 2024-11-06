from langchain_community.document_loaders import PyPDFLoader
loader = PyPDFLoader("docs/202406_PULSO_FGDs ACNUR.pdf")
docs = loader.load_and_split()


from langchain.text_splitter import RecursiveCharacterTextSplitter 
splitter_recursivecharactertext = RecursiveCharacterTextSplitter(
    chunk_size=1000, 
    chunk_overlap=200,
    add_start_index=True,
    separators=["\n\n", "\n", ".", " ", ""],
)
chunks_recursivecharactertext = splitter_recursivecharactertext.split_documents(docs)

#from langchain_community.vectorstores import Chroma
from langchain_community.vectorstores import SKLearnVectorStore

#import chromadb 
#chroma_client = chromadb.PersistentClient(path="persist/")
## A collection is created with the following
#chroma_client.delete_collection("quickstart") # If re-indexing
#chroma_collection = chroma_client.create_collection("quickstart")

from langchain_community.embeddings import GPT4AllEmbeddings 
model_name_embed = "all-MiniLM-L6-v2.gguf2.f16.gguf"
gpt4all_kwargs = {'allow_download': 'True'} #For Edouard: This part needed in tutorial
embeddings_bert = GPT4AllEmbeddings(
    model_name = model_name_embed,
    gpt4all_kwargs=gpt4all_kwargs
)


vectorstore_recursivecharactertext_bert = SKLearnVectorStore.from_documents(
    documents=chunks_recursivecharactertext,
    embedding=embeddings_bert,
#    collection_name= "recursivecharactertext_bert",
    persist_path = "persist")

from langchain_ollama import ChatOllama
ollama_mixtral = ChatOllama(
    model="mistral:latest",  
    temperature=0.2, 
    request_timeout=500
)


RAG_prompt = """
<s> 
[INST]Act if you were a public program evaluation expert working for UNHCR. 
Your audience target is composed of Senior Executives that are managing the operation or program that got evaluated.[/INST]

Your task is to generate an executive summary of the report you just ingested. 
</s>

[INST]
The summary should follow the following defined structure:
 
 - In the first part titled "Access to information", start with a description of information avaialable and include as 5 bullet points, the main sources of information used besides WhatsApp and Facebook, what people have used information sources for, if they consider the information to be useful and reliable and what information they would have liked to have received on their trip. 
 For each point explain what the issues have been and give a practical example from the report.
 
 - In a second part titled: "Travel information" try to review questions related to safety and security, routes and transportation means taken, climate factors and describe at least five needs in detail.  
 
 - In a third part, title: "Sociopolitical Context", describe what people mentioned about the presence of armed or security forces at the border, expulsion matters, legislative changes such as refugee law in Chile, the political situation in Ecuador and their perspectives on requesting asylum.
 
 - In the fourth part: "Perspectives for the Future", include and summarize the main points mentioned. Classify the recommendations according to these categories:
      
      1. "Personal expectations": i.e what people expect or aspire to once they arrive to their destination;
       
      2. "Support networks": i.e whether the persons have family, friends or acquaintances at their destination;
    
      3. "Worries regarding returns": i.e whether they have any concerns if they had to return to their country of origin.
 
 - At the end, for the "Conclusion", craft a reflective conclusion in one sentence that highlights the broader significance of the discussed topic. 
[/INST]
"""

from langchain_core.prompts import ChatPromptTemplate
prompt_retrieval = ChatPromptTemplate.from_template(
"""Answer the following question based only on the provided context:
<context>
{context}
</context>
Question: {input}"""
)

ragRetriever_recursivecharactertext_bert = vectorstore_recursivecharactertext_bert.as_retriever()

from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain

combine_docs_chain_mixtral = create_stuff_documents_chain(
    ollama_mixtral ,
    prompt_retrieval
)
qa_chain_recursivecharactertext_bert = create_retrieval_chain(
    ragRetriever_recursivecharactertext_bert, 
    combine_docs_chain_mixtral
)

response_recursivecharactertext_bert = qa_chain_recursivecharactertext_bert.invoke({"input": RAG_prompt}) 

import docx
from markdown import markdown
import re

def create_word_doc(text, file_name):
    # Create a document
    doc = docx.Document()
    # add a heading of level 0 (largest heading)
    doc.add_heading('FGDs Summary', 0) 

     # Split the text into lines
    lines = text.split('\n')
    # Create a set to store bolded and italic strings
    bolded_and_italic = set()
    for line in lines:
        # Check if the line is a heading
        if line.startswith('#'):
            level = line.count('#')
            doc.add_heading(line[level:].strip(), level)
        else:
            # Check if the line contains markdown syntax for bold or italic
            if '**' in line or '*' in line:
                # Split the line into parts
                parts = re.split(r'(\*{1,2}(.*?)\*{1,2})', line)
                # Add another paragraph
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    # Remove the markdown syntax
                    content = part.strip('*')
                    # Check if the content has been added before
                    if content not in bolded_and_italic:
                        # Add a run with the part and format it
                        run = p.add_run(content)
                        run.font.name = 'Arial'
                        run.font.size = docx.shared.Pt(12)
                        # If the part was surrounded by **, make it bold
                        if '**' in part:
                            run.bold = True
                        # If the part was surrounded by *, make it italic
                        elif '*' in part:
                            run.italic = True
                        # Add the content to the set
                        bolded_and_italic.add(content)
            else:
                # Add another paragraph
                p = doc.add_paragraph()
                # Add a run with the line and format it
                run = p.add_run(line)
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(12)

    ## Add  a disclaimer... ----------------
    # add a page break to start a new page
    doc.add_page_break()
    # add a heading of level 2
    doc.add_heading('DISCLAIMER:', 2)
    doc_para = doc.add_paragraph() 
    doc_para.add_run('This document contains material generated by artificial intelligence technology. While efforts have been made to ensure accuracy, please be aware that AI-generated content may not always fully represent the intent or expertise of human-authored material and may contain errors or inaccuracies. An AI model might generate content that sounds plausible but that is either factually incorrect or unrelated to the given context. These unexpected outcomes, also called AI hallucinations, can stem from biases, under-performing information retrieval, lack of real-world understanding, or limitations in training data.').italic = True

    # Save the document ---------------
    doc.save(file_name)

create_word_doc(
    response_recursivecharactertext_bert["answer"], 
    "FGDs_summary_response_recursivecharactertext_bert.docx"
)