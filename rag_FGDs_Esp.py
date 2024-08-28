import os
from langchain_community.document_loaders.directory import DirectoryLoader
from langchain_community.document_loaders import PyPDFLoader

from langchain_community.vectorstores import Chroma
import chromadb 

from langchain.text_splitter import RecursiveCharacterTextSplitter 

from langchain_community.embeddings import GPT4AllEmbeddings 

from langchain_community.chat_models import ChatOllama

pdf_folder_path = "reports"

docs = []
for file in os.listdir(pdf_folder_path):
        if file.endswith('.pdf'):
            pdf_path = os.path.join(pdf_folder_path, file)
            loader = PyPDFLoader(pdf_path)
            docs.extend(loader.load())

splitter_recursivecharactertext = RecursiveCharacterTextSplitter(
    chunk_size=1000, 
    chunk_overlap=200,
    add_start_index=True,
    separators=["\n\n", "\n", ".", " ", ""],
)
chunks_recursivecharactertext = splitter_recursivecharactertext.split_documents(docs)


chroma_client = chromadb.PersistentClient(path="persist/")
## A collection is created with the following
chroma_client.delete_collection("quickstart") # If re-indexing
chroma_collection = chroma_client.create_collection("quickstart")


model_name_embed = "all-MiniLM-L6-v2.gguf2.f16.gguf"
gpt4all_kwargs = {'allow_download': 'True'} #For Edouard: This part needed in tutorial
embeddings_bert = GPT4AllEmbeddings(
    model_name = model_name_embed,
    gpt4all_kwargs=gpt4all_kwargs
)

vectorstore_recursivecharactertext_bert = Chroma.from_documents(
    documents=chunks_recursivecharactertext,
    embedding=embeddings_bert,
    collection_name= "recursivecharactertext_bert",
    persist_directory = "persist")   


from langchain_community.chat_models import ChatOllama
ollama_mixtral = ChatOllama(
    model="mixtral:8x7b",  
    temperature=0.2, 
    request_timeout=500
)


RAG_prompt = """
<s> 
[INST] Actua como si fueras un experto senior de análisis de programas públicos para ACNUR.
Tu población objetivo es compuesta de Managers Ejecutivos Senior que manejan la operación o el programa para el cual se recolectó la información.
[/INST]

Tu tarea es generar un resumen ejecutivo del reporte que has procesado.
</s>

[INST]
El resumen tendría que seguir la estructura siguiente:

 
    - En la primera parte titulada "Ruta de viaje", describe las 4 principales razones por las que personas refugiadas y migrantes salen de su lugar de origen, con ejemplos concretos y diferenciando por lugar de origen (por ejemplo, Venezuela, Ecuador o Colombia), además describe los lugares que transitaron las personas y las razones para dejar estos lugares, también incluye los medios de transporte utilizados en el viaje y los riesgos que estos medios de transporte pueden suponer.

    - En la segunda parte titulada "Regularización y sistema de asilo" describe que información tenían las personas sobre procesos de regularización antes de emprender su viaje, los conocimientos sobre algún solicitante de asilo en la ruta, y define las tres principales limitaciones para llevar a cabo estos procesos de regularización o pedir asilo. Proporciona ejemplos claros del texto.

    - En la tercera parte titulada "Acceso a la información" incluye las principales fuentes de información utilizadas y para que se ha utilizado cada una de ellas, cual es la fuente de información más fiable y la menos fiable, la utilidad de la información en el viaje y que información les hubiera gustado tener antes de viajar.

    - En la cuarta parte titulada "Contexto sociopolítico" sepárala por los siguientes temas: presencia de fuerzas armadas o de seguridad en la frontera, las expulsiones (PASEE), la ley de asilo en Chile y situación en Ecuador. Y define si cada uno de estos cuatro eventos  las personas los conocían, de qué manera han influido en la toma de decisiones de las personas para emprender su viaje y si han tenido algún impacto en su ruta de viaje. Al final haz un resumen de una frase sobre lo que la gente ha mencionado sobre los cuatro eventos mencionados.

    - En la quinta parte titulada "Percepción de seguridad" describe si las personas se han sentido seguras durante su ruta y que situaciones les han hecho sentir en riesgo, incluyendo algunos grupos que puedan ser especialmente afectados por estas situaciones y proporcionando ejemplos concretos. Además, determina si las personas conocen lugares para recibir protección en caso de que se sientan en riesgo o sufran algún incidente.

    - En la sexta parte titulada "Necesidades" sintetiza las tres principales necesidades en la ruta, con ejemplos concretos, incluye de qué manera los factores climatológicos han afectado estás necesidades de las personas, y las tres principales estrategias que las personas han adaptado para afrontar sus necesidades básicas.

    - En la séptima parte titulada "Perspectivas de futuro" incluye las aspiraciones de las personas para el futuro, si las personas tienen redes de apoyo o familiares en su destino y como esta presencia de redes de apoyo puede influir su decisión de ir a este destino, y los posibles miedos que tienen en caso de que vuelvan a su lugar de origen, dando ejemplos concretos sobre historias personales.

    Al final, para las "Conclusiones" escribe una frase reflectiva que resuma todos los puntos tratados. 
 
[/INST]
"""

from langchain_core.prompts import ChatPromptTemplate
prompt_retrieval = ChatPromptTemplate.from_template(
"""Answer the following question based only on the provided context:
<context>
{context}
</context>
Question: {input}"""
)

ragRetriever_recursivecharactertext_bert = vectorstore_recursivecharactertext_bert.as_retriever()

from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain

combine_docs_chain_mixtral = create_stuff_documents_chain(
    ollama_mixtral ,
    prompt_retrieval
)
qa_chain_recursivecharactertext_bert = create_retrieval_chain(
    ragRetriever_recursivecharactertext_bert, 
    combine_docs_chain_mixtral
)

response_recursivecharactertext_bert = qa_chain_recursivecharactertext_bert.invoke({"input": RAG_prompt}) 

import docx
from markdown import markdown
import re

def create_word_doc(text, file_name):
    # Create a document
    doc = docx.Document()
    # add a heading of level 0 (largest heading)
    doc.add_heading('FGDs Summary', 0) 

     # Split the text into lines
    lines = text.split('\n')
    # Create a set to store bolded and italic strings
    bolded_and_italic = set()
    for line in lines:
        # Check if the line is a heading
        if line.startswith('#'):
            level = line.count('#')
            doc.add_heading(line[level:].strip(), level)
        else:
            # Check if the line contains markdown syntax for bold or italic
            if '**' in line or '*' in line:
                # Split the line into parts
                parts = re.split(r'(\*{1,2}(.*?)\*{1,2})', line)
                # Add another paragraph
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    # Remove the markdown syntax
                    content = part.strip('*')
                    # Check if the content has been added before
                    if content not in bolded_and_italic:
                        # Add a run with the part and format it
                        run = p.add_run(content)
                        run.font.name = 'Arial'
                        run.font.size = docx.shared.Pt(12)
                        # If the part was surrounded by **, make it bold
                        if '**' in part:
                            run.bold = True
                        # If the part was surrounded by *, make it italic
                        elif '*' in part:
                            run.italic = True
                        # Add the content to the set
                        bolded_and_italic.add(content)
            else:
                # Add another paragraph
                p = doc.add_paragraph()
                # Add a run with the line and format it
                run = p.add_run(line)
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(12)

    ## Add  a disclaimer... ----------------
    # add a page break to start a new page
    doc.add_page_break()
    # add a heading of level 2
    doc.add_heading('DISCLAIMER:', 2)
    doc_para = doc.add_paragraph() 
    doc_para.add_run('This document contains material generated by artificial intelligence technology. While efforts have been made to ensure accuracy, please be aware that AI-generated content may not always fully represent the intent or expertise of human-authored material and may contain errors or inaccuracies. An AI model might generate content that sounds plausible but that is either factually incorrect or unrelated to the given context. These unexpected outcomes, also called AI hallucinations, can stem from biases, under-performing information retrieval, lack of real-world understanding, or limitations in training data.').italic = True

    # Save the document ---------------
    doc.save(file_name)

create_word_doc(
    response_recursivecharactertext_bert["answer"], 
    "FGDs_ESP_summary_response_recursivecharactertext_bert.docx"
)