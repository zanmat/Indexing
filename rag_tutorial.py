from langchain_community.document_loaders import PyPDFLoader
loader = PyPDFLoader("5dd4f7d24.pdf")
docs = loader.load_and_split()


from langchain.text_splitter import RecursiveCharacterTextSplitter 
splitter_recursivecharactertext = RecursiveCharacterTextSplitter(
    chunk_size=1000, 
    chunk_overlap=200,
    add_start_index=True,
    separators=["\n\n", "\n", ".", " ", ""],
)
chunks_recursivecharactertext = splitter_recursivecharactertext.split_documents(docs)

from langchain_community.vectorstores import Chroma
import chromadb 
chroma_client = chromadb.PersistentClient(path="persist/")
## A collection is created with the following
# chroma_collection = chroma_client.create_collection('collection')

from langchain_community.embeddings import GPT4AllEmbeddings 
model_name_embed = "all-MiniLM-L6-v2.gguf2.f16.gguf"
gpt4all_kwargs = {'allow_download': 'True'} #For Edouard: This part needed in tutorial
embeddings_bert = GPT4AllEmbeddings(
    model_name = model_name_embed,
    gpt4all_kwargs=gpt4all_kwargs
)

vectorstore_recursivecharactertext_bert = Chroma.from_documents(
    documents=chunks_recursivecharactertext,
    embedding=embeddings_bert,
    collection_name= "recursivecharactertext_bert",
    persist_directory = "persist")

from langchain_community.chat_models import ChatOllama
ollama_mixtral = ChatOllama(
    model="mixtral:8x7b",  
    temperature=0.2, 
    request_timeout=500
)


RAG_prompt = """
<s> 
[INST]Act if you were a public program evaluation expert working for UNHCR. 
Your audience target is composed of Senior Executives that are managing the operation or program that got evaluated.[/INST]

Your task is to generate an executive summary of the report you just ingested. 
</s>

[INST]
The summary should follow the following defined structure:
 
 - In the first part titled "What have we learn?", start with a description of the Forcibly Displaced population in the operation and include as 5 bullet points, the main challenges in relation with the evaluation objectives that have been identified in the document. 
 For each challenge explain why it's a problem and give a practical example to illustrate the consequence of this problem.
 
 - In a second part titled: "How did we get there?" try to review the common root causes for all the challenges that have been identified.  
 
 - In a third part, title: "What is working well?", provide a summary of the main success and achievement, i.e. things that have been identified as good practices and / or effective by the evaluators.
 
 - In the fourth part: "Now What to do?", include and summarize the recommendations proposed by the evaluation. Classify the recommendations according to their relevant level:
      
      1. "Operational Level": i.e recommendations that need to be implemented in the field as an adaptation or change of current practices. Please flag clearly, if this is the case, the recommendations related to practice that should be stopped or discontinued;
       
      2. "Organizational level": i.e recommendations that require changes in staffing or capacity building. Please flag clearly, if this is the case, the recommendations related to practice that should be stopped or discontinued;
    
      3. "Strategic Level": i.e recommendations that require a change in existing policy and rules.
 
 - At the end, for the "Conclusion", craft a reflective conclusion in one sentence that highlights the broader significance of the discussed topic. 
[/INST]
"""

from langchain_core.prompts import ChatPromptTemplate
prompt_retrieval = ChatPromptTemplate.from_template(
"""Answer the following question based only on the provided context:
<context>
{context}
</context>
Question: {input}"""
)

ragRetriever_recursivecharactertext_bert = vectorstore_recursivecharactertext_bert.as_retriever()

from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain

combine_docs_chain_mixtral = create_stuff_documents_chain(
    ollama_mixtral ,
    prompt_retrieval
)
qa_chain_recursivecharactertext_bert = create_retrieval_chain(
    ragRetriever_recursivecharactertext_bert, 
    combine_docs_chain_mixtral
)

response_recursivecharactertext_bert = qa_chain_recursivecharactertext_bert.invoke({"input": RAG_prompt}) 

import docx
from markdown import markdown
import re

def create_word_doc(text, file_name):
    # Create a document
    doc = docx.Document()
    # add a heading of level 0 (largest heading)
    doc.add_heading('Evaluation Brief', 0) 

     # Split the text into lines
    lines = text.split('\n')
    # Create a set to store bolded and italic strings
    bolded_and_italic = set()
    for line in lines:
        # Check if the line is a heading
        if line.startswith('#'):
            level = line.count('#')
            doc.add_heading(line[level:].strip(), level)
        else:
            # Check if the line contains markdown syntax for bold or italic
            if '**' in line or '*' in line:
                # Split the line into parts
                parts = re.split(r'(\*{1,2}(.*?)\*{1,2})', line)
                # Add another paragraph
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    # Remove the markdown syntax
                    content = part.strip('*')
                    # Check if the content has been added before
                    if content not in bolded_and_italic:
                        # Add a run with the part and format it
                        run = p.add_run(content)
                        run.font.name = 'Arial'
                        run.font.size = docx.shared.Pt(12)
                        # If the part was surrounded by **, make it bold
                        if '**' in part:
                            run.bold = True
                        # If the part was surrounded by *, make it italic
                        elif '*' in part:
                            run.italic = True
                        # Add the content to the set
                        bolded_and_italic.add(content)
            else:
                # Add another paragraph
                p = doc.add_paragraph()
                # Add a run with the line and format it
                run = p.add_run(line)
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(12)

    ## Add  a disclaimer... ----------------
    # add a page break to start a new page
    doc.add_page_break()
    # add a heading of level 2
    doc.add_heading('DISCLAIMER:', 2)
    doc_para = doc.add_paragraph() 
    doc_para.add_run('This document contains material generated by artificial intelligence technology. While efforts have been made to ensure accuracy, please be aware that AI-generated content may not always fully represent the intent or expertise of human-authored material and may contain errors or inaccuracies. An AI model might generate content that sounds plausible but that is either factually incorrect or unrelated to the given context. These unexpected outcomes, also called AI hallucinations, can stem from biases, under-performing information retrieval, lack of real-world understanding, or limitations in training data.').italic = True

    # Save the document ---------------
    doc.save(file_name)

create_word_doc(
    response_recursivecharactertext_bert["answer"], 
    "Evaluation_Brief_response_recursivecharactertext_bert.docx"
)