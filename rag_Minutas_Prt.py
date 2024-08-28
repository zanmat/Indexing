import os
from langchain_community.document_loaders.directory import DirectoryLoader
from langchain_community.document_loaders import PyPDFLoader

from langchain_community.vectorstores import Chroma
import chromadb 

from langchain.text_splitter import RecursiveCharacterTextSplitter 

from langchain_community.embeddings import GPT4AllEmbeddings 

from langchain_community.chat_models import ChatOllama

pdf_folder_path = "minutes"

docs = []
for file in os.listdir(pdf_folder_path):
        if file.endswith('.pdf'):
            pdf_path = os.path.join(pdf_folder_path, file)
            loader = PyPDFLoader(pdf_path)
            docs.extend(loader.load())

splitter_recursivecharactertext = RecursiveCharacterTextSplitter(
    chunk_size=1000, 
    chunk_overlap=200,
    add_start_index=True,
    separators=["\n\n", "\n", ".", " ", ""],
)
chunks_recursivecharactertext = splitter_recursivecharactertext.split_documents(docs)


chroma_client = chromadb.PersistentClient(path="persist/")
## A collection is created with the following
chroma_client.delete_collection("quickstart") # If re-indexing
chroma_collection = chroma_client.create_collection("quickstart")


model_name_embed = "all-MiniLM-L6-v2.gguf2.f16.gguf"
gpt4all_kwargs = {'allow_download': 'True'} #For Edouard: This part needed in tutorial
embeddings_bert = GPT4AllEmbeddings(
    model_name = model_name_embed,
    gpt4all_kwargs=gpt4all_kwargs
)

vectorstore_recursivecharactertext_bert = Chroma.from_documents(
    documents=chunks_recursivecharactertext,
    embedding=embeddings_bert,
    collection_name= "recursivecharactertext_bert",
    persist_directory = "persist")   


from langchain_community.chat_models import ChatOllama
ollama_mixtral = ChatOllama(
    model="mixtral:8x7b",  
    temperature=0.2, 
    request_timeout=500
)


RAG_prompt = """
<s> 
[INST] Actua como si fueras un experto senior de análisis de programas públicos para ACNUR encargado de tomar las minutas oficiales de la reunión.
Tu población objetivo es compuesta de Managers Ejecutivos Senior que manejan la operación o el programa y que leerán las minutas.
[/INST]

Tu tarea es generar las minutas de la reunión desde la transcripción que has procesado. Corrige eventuales errores en el texto debidos a la transcripción. Escribe en español.
</s>

[INST]
Las minutas tendrían que seguir la estructura siguiente:

 
    - En la primera parte titulada "Acuerdos", describe los puntos principales acordados durante toda la reunión.

    - En la segunda parte titulada "Reporte de los subsectores" describe los puntos mencionados durante la reunión desde los coordinadores sectoriales. Organiza la respuesta en los siguientes subsectores:
        1- "Acceso a territorio, sistema de asilo y documentación - Encuentros SJM & ACNUR" para los temas relacionados con regularización, acceso al territorio, sistema de refugio y acceso a documentación;
        2- "Protección de Niñez y Adolescencia – UNICEF& Aldeas Infantiles" para los temas relacionados con niñas, niños y adolescentes, sus derechos y desafíos mencionados;
        3- "Violencia Basada en Género – UNFPA & ACNUR" para los puntos mencionados relacionados con Violencia Basada en Género, desde los coordinadores del subsector;
        4- "Trata y Tráfico – OIM & UNODC" para todos los temas relacionados con trata de personas y tráfico ilícito de migrantes.

    - En la tercera parte titulada "Presentación de nuevos miembros" incluye en bullet points separados un resumen de cada nuevo miembro introducido durante la reunión y su intervención.

 
[/INST]
"""

from langchain_core.prompts import ChatPromptTemplate
prompt_retrieval = ChatPromptTemplate.from_template(
"""Answer the following question based only on the provided context:
<context>
{context}
</context>
Question: {input}"""
)

ragRetriever_recursivecharactertext_bert = vectorstore_recursivecharactertext_bert.as_retriever()

from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain

combine_docs_chain_mixtral = create_stuff_documents_chain(
    ollama_mixtral ,
    prompt_retrieval
)
qa_chain_recursivecharactertext_bert = create_retrieval_chain(
    ragRetriever_recursivecharactertext_bert, 
    combine_docs_chain_mixtral
)

response_recursivecharactertext_bert = qa_chain_recursivecharactertext_bert.invoke({"input": RAG_prompt}) 

import docx
from markdown import markdown
import re

def create_word_doc(text, file_name):
    # Create a document
    doc = docx.Document()
    # add a heading of level 0 (largest heading)
    doc.add_heading('Minutas Subgrupo de Protección', 0) 

     # Split the text into lines
    lines = text.split('\n')
    # Create a set to store bolded and italic strings
    bolded_and_italic = set()
    for line in lines:
        # Check if the line is a heading
        if line.startswith('#'):
            level = line.count('#')
            doc.add_heading(line[level:].strip(), level)
        else:
            # Check if the line contains markdown syntax for bold or italic
            if '**' in line or '*' in line:
                # Split the line into parts
                parts = re.split(r'(\*{1,2}(.*?)\*{1,2})', line)
                # Add another paragraph
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    # Remove the markdown syntax
                    content = part.strip('*')
                    # Check if the content has been added before
                    if content not in bolded_and_italic:
                        # Add a run with the part and format it
                        run = p.add_run(content)
                        run.font.name = 'Arial'
                        run.font.size = docx.shared.Pt(12)
                        # If the part was surrounded by **, make it bold
                        if '**' in part:
                            run.bold = True
                        # If the part was surrounded by *, make it italic
                        elif '*' in part:
                            run.italic = True
                        # Add the content to the set
                        bolded_and_italic.add(content)
            else:
                # Add another paragraph
                p = doc.add_paragraph()
                # Add a run with the line and format it
                run = p.add_run(line)
                run.font.name = 'Arial'
                run.font.size = docx.shared.Pt(12)

    ## Add  a disclaimer... ----------------
    # add a page break to start a new page
    doc.add_page_break()
    # add a heading of level 2
    doc.add_heading('DISCLAIMER:', 2)
    doc_para = doc.add_paragraph() 
    doc_para.add_run('This document contains material generated by artificial intelligence technology. While efforts have been made to ensure accuracy, please be aware that AI-generated content may not always fully represent the intent or expertise of human-authored material and may contain errors or inaccuracies. An AI model might generate content that sounds plausible but that is either factually incorrect or unrelated to the given context. These unexpected outcomes, also called AI hallucinations, can stem from biases, under-performing information retrieval, lack of real-world understanding, or limitations in training data.').italic = True

    # Save the document ---------------
    doc.save(file_name)

create_word_doc(
    response_recursivecharactertext_bert["answer"], 
    "Minutas_Prt_ESP_summary_response_recursivecharactertext_bert.docx"
)